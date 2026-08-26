"""Minimal Markdown -> .docx renderer for the labgate technical documents.

Supported: # ## ### headings · paragraphs · - bullets · 1. numbered ·
tables · ```code``` · ![cap](path){w} figures · > callouts · --- rules ·
inline **bold** and `code`.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK   = RGBColor(0x22, 0x22, 0x22)
BLUE  = RGBColor(0x2F, 0x5A, 0x93)
GREEN = RGBColor(0x2E, 0x8B, 0x57)
MUTE  = RGBColor(0x66, 0x66, 0x66)
CODEC = RGBColor(0x1E, 0x46, 0x20)


def _shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def _emit(par, text, size, color, bold=False, italic=False):
    """Emit runs for `text`, honouring `code` spans, with inherited style."""
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("`") and part.endswith("`") and len(part) > 2:
            r = par.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(size - 0.7)
            r.font.color.rgb = CODEC
        else:
            r = par.add_run(part)
            r.font.size = Pt(size)
            if color is not None:
                r.font.color.rgb = color
        r.bold = bold
        r.italic = italic


def _inline(par, text, size=10.5, color=INK, italic=False):
    """Render **bold**, *italic* and `code`, including code nested in bold."""
    for part in re.split(r"(\*\*.+?\*\*|\*[^*\n]+?\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            _emit(par, part[2:-2], size, color, bold=True, italic=italic)
        elif (part.startswith("*") and part.endswith("*")
              and not part.startswith("**") and len(part) > 2):
            _emit(par, part[1:-1], size, color, italic=True)
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            _emit(par, part, size, color, italic=italic)
        else:
            _emit(par, part, size, color, italic=italic)


def build(md_path, out_path, title, subtitle, meta_line, fig_dir):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10.5)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.14
    for name, size, color in [("Heading 1", 15.5, BLUE), ("Heading 2", 12.5, BLUE),
                              ("Heading 3", 11, GREEN)]:
        s = doc.styles[name]
        s.font.name = "Calibri"; s.font.size = Pt(size)
        s.font.color.rgb = color; s.font.bold = True
        s.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        s.paragraph_format.space_after = Pt(4)

    # ---- title block
    p = doc.add_paragraph(); r = p.add_run(title)
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph(); r = p.add_run(subtitle)
    r.font.size = Pt(12.5); r.font.color.rgb = BLUE
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph(); r = p.add_run(meta_line)
    r.font.size = Pt(9.5); r.font.color.rgb = MUTE

    lines = Path(md_path).read_text().splitlines()
    i, fig_n = 0, 0
    while i < len(lines):
        line = lines[i]

        if not line.strip():
            i += 1; continue

        # ---- horizontal rule / page break
        if line.strip() == "---":
            doc.add_page_break(); i += 1; continue

        # ---- headings
        m = re.match(r"^(#{1,3})\s+(.*)", line)
        if m:
            doc.add_heading(m.group(2).strip(), level=len(m.group(1)))
            i += 1; continue

        # ---- figure
        m = re.match(r"^!\[(.*?)\]\((.*?)\)(?:\{(\d+(?:\.\d+)?)\})?", line)
        if m:
            cap, path, w = m.group(1), m.group(2), m.group(3)
            fig_n += 1
            doc.add_picture(str(Path(fig_dir) / Path(path).name),
                            width=Inches(float(w) if w else 6.3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cap:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f"Figure {fig_n} — {cap}")
                r.italic = True; r.font.size = Pt(9); r.font.color.rgb = MUTE
                p.paragraph_format.space_after = Pt(10)
            i += 1; continue

        # ---- code block
        if line.startswith("```"):
            i += 1; buf = []
            while i < len(lines) and not lines[i].startswith("```"):
                buf.append(lines[i]); i += 1
            i += 1
            t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
            cell = t.cell(0, 0); _shade(cell, "F6F6F4")
            cell.paragraphs[0].text = ""
            r = cell.paragraphs[0].add_run("\n".join(buf))
            r.font.name = "Consolas"; r.font.size = Pt(8.4)
            r.font.color.rgb = RGBColor(0x2A, 0x2A, 0x2A)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # ---- callout
        if line.startswith("> "):
            buf = []
            while i < len(lines) and lines[i].startswith("> "):
                buf.append(lines[i][2:]); i += 1
            t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
            cell = t.cell(0, 0); _shade(cell, "EAF1FA")
            cell.paragraphs[0].text = ""
            _inline(cell.paragraphs[0], " ".join(buf), size=9.8, color=None, italic=True)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # ---- table
        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=len(rows), cols=ncol)
            t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
            t.autofit = True
            for ri, row in enumerate(rows):
                trPr = t.rows[ri]._tr.get_or_add_trPr()
                cant = OxmlElement("w:cantSplit"); trPr.append(cant)
                if ri == 0:
                    hdr = OxmlElement("w:tblHeader"); trPr.append(hdr)
                for ci in range(ncol):
                    cell = t.rows[ri].cells[ci]
                    cell.text = ""
                    txt = row[ci] if ci < len(row) else ""
                    _inline(cell.paragraphs[0], txt, size=9,
                            color=RGBColor(0xFF, 0xFF, 0xFF) if ri == 0 else INK)
                    if ri == 0:
                        for rr in cell.paragraphs[0].runs:
                            rr.bold = True
                        _shade(cell, "2F5A93")
                    elif ri % 2 == 0:
                        _shade(cell, "F2F5FA")
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        # ---- bullets / numbered
        m = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)", line)
        if m:
            indent, marker, text = m.group(1), m.group(2), m.group(3)
            style = "List Number" if marker[0].isdigit() else "List Bullet"
            p = doc.add_paragraph(style=style)
            if indent:
                p.paragraph_format.left_indent = Inches(0.55)
            p.paragraph_format.space_after = Pt(3)
            _inline(p, text)
            i += 1; continue

        # ---- paragraph (join wrapped lines)
        buf = [line]
        i += 1
        while (i < len(lines) and lines[i].strip()
               and not re.match(r"^(#{1,3}\s|!\[|\||```|>\s|\s*[-*]\s|\s*\d+\.\s|---$)", lines[i])):
            buf.append(lines[i]); i += 1
        p = doc.add_paragraph()
        _inline(p, " ".join(x.strip() for x in buf))

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    md, out, title, subtitle, meta, figs = sys.argv[1:7]
    print("saved:", build(md, out, title, subtitle, meta, figs))
